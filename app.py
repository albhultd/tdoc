import time
import streamlit as st
import pandas as pd
import os
from dotenv import load_dotenv
import search  # Import the search module
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

load_dotenv()

# Initialize counters
user_counter = 0
question_counter = 0

st.set_page_config(
        page_title="DocGPT GT",
        page_icon="speech_balloon",
        layout="wide",
    )

hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            footer:after {
	content:'2023'; 
	visibility: visible;
	display: block;
	position: relative;
	padding: 5px;
	top: 2px;
}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

def save_as_pdf(conversation):
    pdf_filename = "conversation.pdf"
    c = canvas.Canvas(pdf_filename, pagesize=letter)
   
    c.drawString(100, 750, "Conversation:")
    y_position = 730
    for q, a in conversation:
        c.drawString(120, y_position, f"Q: {q}")
        c.drawString(120, y_position - 20, f"A: {a}")
        y_position -= 40
   
    c.save()
   
    st.markdown(f"Download [PDF](./{pdf_filename})")

def save_as_docx(conversation):
    doc = Document()
    doc.add_heading('Conversation', 0)
   
    for q, a in conversation:
        doc.add_paragraph(f'Q: {q}')
        doc.add_paragraph(f'A: {a}')
   
    doc_filename = "conversation.docx"
    doc.save(doc_filename)
   
    st.markdown(f"Download [DOCX](./{doc_filename})")

def save_as_xlsx(conversation):
    df = pd.DataFrame(conversation, columns=["Question", "Answer"])
    xlsx_filename = "conversation.xlsx"
    df.to_excel(xlsx_filename, index=False)
   
    st.markdown(f"Download [XLSX](./{xlsx_filename})")

def save_as_txt(conversation):
    txt_filename = "conversation.txt"
    with open(txt_filename, "w") as txt_file:
        for q, a in conversation:
            txt_file.write(f"Q: {q}\nA: {a}\n\n")
   
    st.markdown(f"Download [TXT](./{txt_filename})")

def main():
    global user_counter, question_counter
    user_counter += 1  # Increment user counter
    st.markdown(f"Number of Users: {user_counter}")

    st.markdown('<h1>Ask anything from texts</h1><p style="font-size: 12; color: gray;"></p>', unsafe_allow_html=True)
    st.markdown("<h2>Upload documents</h2>", unsafe_allow_html=True)
    uploaded_files = st.file_uploader("Upload one or more documents", type=['pdf', 'docx'], accept_multiple_files=True)
    question = st.text_input("Ask a question based on the documents", key="question_input")
    question_counter += 1  # Increment question counter
    st.markdown(f"Number of Questions: {question_counter}")

    progress = st.progress(0)
    for i in range(100):
        progress.progress(i + 1)
        time.sleep(0.01)

    if uploaded_files:
        df = pd.DataFrame(columns=["page_num", "paragraph_num", "content", "tokens"])
        for uploaded_file in uploaded_files:
            paragraphs = search.read_pdf_pdfminer(uploaded_file) if uploaded_file.type == "application/pdf" else search.read_docx(uploaded_file)
            temp_df = pd.DataFrame(
                [(p.page_num, p.paragraph_num, p.content, search.count_tokens(p.content, search.tokenizer))
                for p in paragraphs],
                columns=["page_num", "paragraph_num", "content", "tokens"]
            )
            df = pd.concat([df, temp_df], ignore_index=True)

        if "interactions" not in st.session_state:
            st.session_state["interactions"] = []

        answer = ""
        if question != st.session_state.get("last_question", ""):
            st.text("Searching...")
            answer = search.answer_query_with_context(question, df)
            st.session_state["interactions"].append((question, answer))
            st.write(answer)

        st.markdown("### Interaction History")
        for q, a in st.session_state["interactions"]:
            st.write(f"**Q:** {q}\n\n**A:** {a}")

        st.session_state["last_question"] = question

        st.markdown("<h2>Sample paragraphs</h2>", unsafe_allow_html=True)
        sample_size = min(len(df), 5)
        st.dataframe(df.sample(n=sample_size))  

        if st.button("Save as PDF"):
            save_as_pdf(st.session_state["interactions"])
        if st.button("Save as DOCX"):
            save_as_docx(st.session_state["interactions"])
        if st.button("Save as XLSX"):
            save_as_xlsx(st.session_state["interactions"])
        if st.button("Save as TXT"):
            save_as_txt(st.session_state["interactions"])


    else:
        st.markdown("<h2>Please upload a document to proceed.</h2>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
